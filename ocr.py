from pdf2image import convert_from_path
import pytesseract
import os
from PIL import Image
import numpy as np
from PIL import Image
import time
from docx import Document
import easyocr
from paddleocr import PaddleOCR
#from paddleocr import PaddleOCR
from ImagePreprocee import *
from paddocr import *

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
#ocr = PaddleOCR(use_textline_orientation=True, lang='en')


def convert_pdf_to_images(PDF_file):
    output_folder="outputdata"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Convert PDF pages to images
    pages = convert_from_path(PDF_file, dpi=300, thread_count=4)
    
    for i, image in enumerate(pages):
        image_path = os.path.join(output_folder, f"page_{i+1}.jpg")
        image.save(image_path, "JPEG")
        print(f"Saved: {image_path}")
        
def pre_process_img():
    folder_path = r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\outputdata"  # Change this to your folder path

    # Get all file paths
    file_paths = []
    i = 0
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):  # Only include files, not subdirectories
            correct_image = correct_image_rotation(file_path)
            image_path = os.path.join(folder_path, f"correct_page_{i+1}.jpg")
            correct_image.save(image_path, "JPEG")
            file_paths.append(image_path)
            i = i+1

    # Print all file paths
    output_folder_1="preprocessImage"
    if not os.path.exists(output_folder_1):
        os.makedirs(output_folder_1)
    
    for j,path in enumerate(file_paths):
        pre_img = preprocess_image(path)
        image_path = os.path.join(output_folder_1, f"img_{j+1}.jpg")
        cv2.imwrite(image_path, pre_img )
        print(image_path)

"""
def extract_text_pdf_using_ocr(file_path):
    output_docx="text_list_output.docx"
    start_time = time.perf_counter()
    convert_pdf_to_images(file_path)
    pre_process_img()
    
    folder_path = r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\preprocessImage"
    text = []
    
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            extracted_text = pytesseract.image_to_string(file_path)
            text.append(extracted_text.strip())
            
     
    end_time = time.perf_counter()
    print(f"\nTotal processing time: {end_time - start_time:.2f} seconds")
    
    folder_path_1 = r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\outputdata"
    if os.path.exists(folder_path_1):
        for filename in os.listdir(folder_path_1):
            file_path = os.path.join(folder_path_1, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
      
    doc = Document()

    for extract in text:
        doc.add_paragraph(extract)
    
    doc.save(output_docx)
    print("\n Document Saved")
"""
def extract_text_pdf_using_paddleocr(file_path):
    #start_time = time.perf_counter()
    convert_pdf_to_images(file_path)
    pre_process_img()
    start_time = time.perf_counter()
    extract_text_paddleocr()
    end_time = time.perf_counter()
    print(f"\nTotal processing time: {end_time - start_time:.2f} seconds")
    #folder_path = r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\outputdata"
    #folder_path_1 = r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\preprocessImage"
"""
    if os.path.exists(folder_path):
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
        print(folder_path+" Successfully Remove")
    else:
        print(folder_path+"Not Visible in the data set")"""
    
"""
    if os.path.exists(folder_path_1):
        for filename in os.listdir(folder_path_1):
            file_path = os.path.join(folder_path_1, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
        print(folder_path_1+" Successfully Remove")
    else:
        print(folder_path_1+"Not Visible in the data set")"""   



"""
def extract_text_pdf_using_easyocr(file_path):
    reader = easyocr.Reader(['en'])
    output_docx="text_list_output.docx"
    start_time = time.perf_counter()
    convert_pdf_to_images(file_path)
    pre_process_img()
    
    folder_path = r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\preprocessImage"
    text_list = []
    
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            img = Image.open(file_path)
            max_width = 1200
            
            if img.width > max_width:
                ratio = max_width / img.width
                img = img.resize((int(img.width * ratio), int(img.height * ratio)))
            

            result = reader.readtext(np.array(img))
            #result = reader.readtext(file_path)
            texts = [text[1] for text in result]
            text_list.append(texts)
            os.remove(file_path)
            
            
     
    end_time = time.perf_counter()
    print(f"\nTotal processing time: {end_time - start_time:.2f} seconds")
    
    folder_path_1 = r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\outputdata"
    if os.path.exists(folder_path_1):
        for filename in os.listdir(folder_path_1):
            file_path = os.path.join(folder_path_1, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
      
    doc = Document()

    for line in text_list:
        #doc.add_paragraph(extract)
        doc.add_paragraph(line)
    
    doc.save(output_docx)
    print("\n Document Saved")"""
            
pdf_file_path =r"D:\Hexabins_AI_ML_Intern\behaviourapp_ocr\Scanned_PDF\Behaviour_ocr_scanned_Pdf\Inputdata\ABC Behaviour Recording Chart_WhiteR_21-06-2023_GC.pdf"
extract_text_pdf_using_paddleocr(pdf_file_path)
#extract_text_pdf_using_easyocr(pdf_file_path)
#extract_text()